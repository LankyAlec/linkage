# -*- coding: utf-8 -*-
import re
import sys
import os
import json
from urllib.parse import quote_plus
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# === Verifica argomenti ===
with open("log_elabora.txt", "w") as log:
    log.write(f"sys.argv: {sys.argv}\n")
    if len(sys.argv) != 4:
        log.write("❌ Numero errato di argomenti. Attesi 3 (input, output, urn_index).\n")
        log.write("Usage: python elabora_docx.py input.docx output.docx urn_index.json\n")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    urn_index_path = sys.argv[3]
    log.write(f">>> Cerco urn_index.json in: {urn_index_path}\n")


if not os.path.isfile(urn_index_path):
    with open("log_elabora.txt", "a") as log:
        log.write(f"❌ File URN index non trovato: {urn_index_path}\n")
    sys.exit(2)

# === Carica mappa URN da JSON ===
with open(urn_index_path, encoding="utf-8") as f:
    urn_db = json.load(f)

sites = {
    'normattiva': 'https://www.normattiva.it',
    'corte_cassazione': 'https://www.cortedicassazione.it',
    'giustizia_amministrativa': 'https://www.giustizia-amministrativa.it',
}

# === Helpers ===
def roman_to_int(s):
    s = s.upper()
    vals = {'I':1,'V':5,'X':10,'L':50,'C':100,'D':500,'M':1000}
    total = 0
    prev = 0
    for ch in reversed(s):
        v = vals.get(ch, 0)
        if v < prev:
            total -= v
        else:
            total += v
            prev = v
    return total

_EXT_TOKEN = r"(bis|ter|quater|quinquies|sexies|septies|octies|novies|decies|undecies|duodecies|terdecies|quaterdecies|quindecies|sexdecies|septiesdecies|duodevicies|undevicies|vices|vicessemel|vicesbis|vicester|vicesquater|vicesquinquies|vicessexies|vicessepties|duodetricies|undetricies|tricies|triciessemel|triciesbis|triciester|triciesquater|triciesquinquies|triciessexies|triciessepties|duodequadragies|undequadragies|quadragies|quadragiessemel|quadragiesbis|quadragiester|quadragiesquater|quadragiesquinquies|quadragiessexies|quadragiessepties|duodequinquagies|undequinquagies)"
_ART_RE = re.compile(
    rf"(?P<num>\d+(?:\.\d+)?)(?:\s*[-–\.]?\s*(?P<ext>{_EXT_TOKEN}(?:\.\d+)?))?"
    rf"|(?P<num2>\d+)(?P<ext2>{_EXT_TOKEN}(?:\.\d+)?)",
    re.IGNORECASE
)

def parse_articolo(chunk):
    m = _ART_RE.fullmatch(chunk.strip())
    if not m:
        return (chunk.strip(), None)
    if m.group('num'):
        art = m.group('num')
        ext = m.group('ext')
    else:
        art = m.group('num2')
        ext = m.group('ext2')
    return (art, ext.lower() if ext else None)

_ALLEG_RE = re.compile(r"\ballegato\s+(\d+)", re.IGNORECASE)

def parse_allegato(text):
    ma = _ALLEG_RE.search(text)
    return ma.group(1) if ma else None

def build_urn(tipo, numero, data="0000-00-00", articolo=None, estensione=None, allegato=None, vig=None, originale=False):
    base = "https://www.normattiva.it/uri-res/N2Ls?urn=nir:stato:"
    urn = base + tipo + f":{data};{numero}"
    if allegato:
        urn += f":{allegato}"
    if articolo:
        suffix = f"~art{articolo}"
        if estensione:
            suffix += estensione.lower()
        if originale:
            suffix += "@originale"
        if vig:
            suffix += f"!vig={vig}"
        urn += suffix
    return urn

def lookup_urn_entry(raw_key):
    normalized = re.sub(r"\.+", ".", raw_key.lower().replace(" ", ""))
    candidates = [normalized]
    if normalized.endswith("."):
        candidates.append(normalized.rstrip("."))
    else:
        candidates.append(normalized + ".")
    for candidate in candidates:
        if candidate in urn_db:
            return urn_db[candidate]
    return None

def build_codice_urn(match):
    entry = lookup_urn_entry(match.group(2))
    if not entry:
        return None
    return build_urn(
        tipo=entry["tipo"],
        numero=entry["numero"],
        data=entry["data"],
        articolo=parse_articolo(match.group(1))[0],
        estensione=parse_articolo(match.group(1))[1],
        allegato=parse_allegato(match.group(0))
    )

# === Regole di conversione ===
rules = [
    (
        re.compile(r'\b(?:art\.?|articolo)\s+([0-9IVXLCDM\.]+(?:'+_EXT_TOKEN+'(?:\.\d+)?)?)\s+(c\.p+\.\?|c\.c\.\?|c\.p\.c\.\?|c\.p\.p\.\?)\b', re.IGNORECASE),
        lambda m: (
            m.group(0),
            build_codice_urn(m)
        )
    ),
    (
        re.compile(r'\b(?:art\.?|articolo)\s+([IVXLCDM]+|\d+)\s+(?:Cost\.?|Costituzione)\b', re.IGNORECASE),
        lambda m: (
            m.group(0),
            build_urn(
                tipo="costituzione",
                numero="",
                data="1947-12-27",
                articolo=str(roman_to_int(m.group(1))) if re.fullmatch(r'[IVXLCDM]+', m.group(1), re.I) else m.group(1)
            ).replace(";", "", 1)
        )
    )
]

# === Utility per i link ===
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def clear_paragraph_runs(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)

def collect_matches(text):
    matches = []
    for pattern, handler in rules:
        for match in pattern.finditer(text):
            label, link = handler(match)
            if not link:
                continue
            matches.append((match.start(), match.end(), label, link))
    matches.sort(key=lambda item: (item[0], -(item[1] - item[0])))
    filtered = []
    last_end = 0
    for start, end, label, link in matches:
        if start < last_end:
            continue
        filtered.append((start, end, label, link))
        last_end = end
    return filtered

# === Main ===
doc = Document(input_path)

for para in doc.paragraphs:
    text = para.text
    matches = collect_matches(text)
    if not matches:
        continue
    clear_paragraph_runs(para)
    last_end = 0
    for start, end, label, link in matches:
        if start > last_end:
            para.add_run(text[last_end:start])
        add_hyperlink(para, link, label)
        last_end = end
    if last_end < len(text):
        para.add_run(text[last_end:])

# Salva
if os.path.exists(output_path):
    os.remove(output_path)
doc.save(output_path)
